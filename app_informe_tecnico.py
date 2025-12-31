import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import io
from datetime import datetime

# Configuración de la página
st.set_page_config(
    page_title="Generador de Informes Técnicos", page_icon="📄", layout="wide"
)

# Título principal
st.title("📄 Generador de Informes Técnicos - Rotomaquinas SAS")
st.markdown("---")


# Función para establecer bordes de celda
def set_cell_border(cell, **kwargs):
    """
    Establecer bordes de celda en la tabla.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Lista de bordes
    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)
            element = OxmlElement(tag)
            for key in edge_data:
                element.set(qn("w:{}".format(key)), str(edge_data[key]))
            tcPr.append(element)


# Función para agregar imagen centrada en celda
def add_image_to_cell(cell, image_bytes, width_inches=2.5):
    """
    Agregar imagen centrada en una celda.
    """
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(image_bytes), width=Inches(width_inches))


import streamlit as st


# Función principal para crear el documento
def crear_documento_tecnico(datos_empresa, datos_cliente, actividades):
    """
    Crear el documento técnico completo.
    """
    doc = Document()

    # Estilos Globales
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)

    # Configurar márgenes (2.54 cm ~ 1 pulgada)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # ============= ENCABEZADO (MEMBRETE) =============
        header = section.header
        header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        header_table.autofit = False

        # Celda Logo (Izquierda)
        cell_logo = header_table.cell(0, 0)
        cell_logo.width = Inches(2.0)
        try:
            # Ruta relativa o absoluta al logo
            logo_path = "assets/logo.png"
            # Verificar si existe, si no, intentar ruta absoluta construida o manejar error
            paragraph = cell_logo.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(1.3))
        except Exception as e:
            cell_logo.text = "[LOGO]"
            print(f"Error cargando logo: {e}")

        # Celda Información Empresa (Derecha)
        cell_info = header_table.cell(0, 1)
        cell_info.width = Inches(4.5)
        paragraph = cell_info.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = paragraph.add_run("ROTOMAQUINAS S.A.S\n")
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 51, 102)  # Azul oscuro corporativo

        run = paragraph.add_run("Servicios Operativos con Máquinas y Personal\n")
        run.font.size = Pt(9)
        run.font.bold = True

        run = paragraph.add_run("Palmira - Valle del Cauca\n")
        run.font.size = Pt(9)

        run = paragraph.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
        run.font.size = Pt(8)
        run.font.italic = True

    # ============= TÍTULO DEL REPORTE =============
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(
        f"INFORME TÉCNICO: {datos_empresa['nombre_proyecto'].upper()}"
    )
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)  # Azul oscuro

    # Espacio
    doc.add_paragraph()

    # ============= TABLA DE DATOS GENERALES (Rediseñada) =============
    # Usaremos una tabla con bordes más sutiles o solo internos
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"  # Mantenemos grid pero podríamos personalizar

    # Datos a llenar
    data_rows = [
        ("FECHA DEL SERVICIO:", datos_empresa["fecha"]),
        ("TÉCNICO RESPONSABLE:", datos_empresa["tecnico"]),
        ("UBICACIÓN:", datos_empresa["ubicacion"]),
        ("ASUNTO:", "Servicio de mantenimiento y limpieza"),
    ]

    for i, (label, value) in enumerate(data_rows):
        row = table.rows[i]

        # Etiqueta
        cell_label = row.cells[0]
        cell_label.width = Inches(2.5)
        p = cell_label.paragraphs[0]
        run = p.add_run(label)
        run.font.bold = True
        run.font.size = Pt(10)
        # Sombreado gris claro para etiquetas
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), "F2F2F2")
        cell_label._element.get_or_add_tcPr().append(shading_elm)

        # Valor
        cell_value = row.cells[1]
        cell_value.text = value
        cell_value.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ============= DATOS DEL CLIENTE =============
    p = doc.add_paragraph()
    run = p.add_run("DATOS DEL CLIENTE")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 51, 102)

    tabla_cliente = doc.add_table(rows=3, cols=2)
    tabla_cliente.style = "Table Grid"

    client_data = [
        ("RAZÓN SOCIAL / NOMBRE:", datos_cliente["nombre"]),
        ("NIT / C.C:", datos_cliente["nit"]),
        ("DIRECCIÓN:", datos_cliente["direccion"]),
    ]

    for i, (label, value) in enumerate(client_data):
        row = tabla_cliente.rows[i]

        cell_label = row.cells[0]
        cell_label.width = Inches(2.5)
        p = cell_label.paragraphs[0]
        run = p.add_run(label)
        run.font.bold = True
        run.font.size = Pt(10)

        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), "F2F2F2")
        cell_label._element.get_or_add_tcPr().append(shading_elm)

        cell_value = row.cells[1]
        cell_value.text = value
        cell_value.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ============= OBJETIVO =============
    p = doc.add_paragraph()
    run = p.add_run("OBJETIVO")
    run.font.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph(datos_empresa["objetivo"])

    doc.add_paragraph()

    # ============= NOTA =============
    p = doc.add_paragraph()
    run = p.add_run("NOTA:  ")
    run.font.bold = True
    run = p.add_run(datos_empresa["nota"])

    doc.add_paragraph()

    # ============= REGISTRO FOTOGRÁFICO =============
    p = doc.add_paragraph()
    run = p.add_run("REGISTRO FOTOGRÁFICO")
    run.font.bold = True
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ============= ACTIVIDADES =============
    for actividad in actividades:
        # Título de la actividad
        p = doc.add_paragraph()
        run = p.add_run(actividad["titulo"].upper())
        run.font.bold = True
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Si solo hay observación (sin ANTES/DESPUÉS)
        if actividad["tipo"] == "solo_observacion":
            # Tabla de observación
            tabla_obs = doc.add_table(rows=2, cols=1)
            tabla_obs.style = "Table Grid"

            # Fila 1: OBSERVACIÓN
            cell = tabla_obs.rows[0].cells[0]
            p_cell = cell.paragraphs[0]
            run = p_cell.add_run("OBSERVACIÓN: ")
            run.font.bold = True
            p_cell.add_run(actividad["observacion"])

            # Aplicar bordes
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "000000"},
                bottom={"sz": 12, "val": "single", "color": "000000"},
                left={"sz": 12, "val": "single", "color": "000000"},
                right={"sz": 12, "val": "single", "color": "000000"},
            )

            # Fila 2: Imágenes
            cell_img = tabla_obs.rows[1].cells[0]
            if actividad.get("imagenes"):
                # Crear una tabla interna para organizar las imágenes
                num_imagenes = len(actividad["imagenes"])
                cols = min(num_imagenes, 2)  # Máximo 2 columnas
                rows = (num_imagenes + cols - 1) // cols

                for idx, img_bytes in enumerate(actividad["imagenes"]):
                    if idx > 0 and idx % cols == 0:
                        cell_img.add_paragraph()
                    add_image_to_cell(cell_img, img_bytes, width_inches=2.2)

            # Aplicar bordes
            set_cell_border(
                cell_img,
                top={"sz": 12, "val": "single", "color": "000000"},
                bottom={"sz": 12, "val": "single", "color": "000000"},
                left={"sz": 12, "val": "single", "color": "000000"},
                right={"sz": 12, "val": "single", "color": "000000"},
            )

        # Si hay ANTES/DESPUÉS
        elif actividad["tipo"] == "antes_despues":
            # ANTES
            if actividad.get("antes"):
                tabla_antes = doc.add_table(rows=3, cols=1)
                tabla_antes.style = "Table Grid"

                # Fila 1: Título ANTES
                cell = tabla_antes.rows[0].cells[0]
                p_cell = cell.paragraphs[0]
                run = p_cell.add_run("ANTES")
                run.font.bold = True
                run.font.size = Pt(11)
                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Color de fondo
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:fill"), "D9D9D9")
                cell._element.get_or_add_tcPr().append(shading_elm)

                set_cell_border(
                    cell,
                    top={"sz": 12, "val": "single", "color": "000000"},
                    bottom={"sz": 12, "val": "single", "color": "000000"},
                    left={"sz": 12, "val": "single", "color": "000000"},
                    right={"sz": 12, "val": "single", "color": "000000"},
                )

                # Fila 2: Observación ANTES
                cell = tabla_antes.rows[1].cells[0]
                p_cell = cell.paragraphs[0]
                run = p_cell.add_run("OBSERVACIÓN: ")
                run.font.bold = True
                p_cell.add_run(actividad["antes"]["observacion"])

                set_cell_border(
                    cell,
                    top={"sz": 12, "val": "single", "color": "000000"},
                    bottom={"sz": 12, "val": "single", "color": "000000"},
                    left={"sz": 12, "val": "single", "color": "000000"},
                    right={"sz": 12, "val": "single", "color": "000000"},
                )

                # Fila 3: Imágenes ANTES
                cell_img = tabla_antes.rows[2].cells[0]
                if actividad["antes"].get("imagenes"):
                    for img_bytes in actividad["antes"]["imagenes"]:
                        add_image_to_cell(cell_img, img_bytes, width_inches=2.2)

                set_cell_border(
                    cell_img,
                    top={"sz": 12, "val": "single", "color": "000000"},
                    bottom={"sz": 12, "val": "single", "color": "000000"},
                    left={"sz": 12, "val": "single", "color": "000000"},
                    right={"sz": 12, "val": "single", "color": "000000"},
                )

                doc.add_paragraph()

            # DESPUÉS
            if actividad.get("despues"):
                tabla_despues = doc.add_table(rows=3, cols=1)
                tabla_despues.style = "Table Grid"

                # Fila 1: Título DESPUÉS
                cell = tabla_despues.rows[0].cells[0]
                p_cell = cell.paragraphs[0]
                run = p_cell.add_run("DESPUÉS")
                run.font.bold = True
                run.font.size = Pt(11)
                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Color de fondo
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:fill"), "D9D9D9")
                cell._element.get_or_add_tcPr().append(shading_elm)

                set_cell_border(
                    cell,
                    top={"sz": 12, "val": "single", "color": "000000"},
                    bottom={"sz": 12, "val": "single", "color": "000000"},
                    left={"sz": 12, "val": "single", "color": "000000"},
                    right={"sz": 12, "val": "single", "color": "000000"},
                )

                # Fila 2: Observación DESPUÉS
                cell = tabla_despues.rows[1].cells[0]
                p_cell = cell.paragraphs[0]
                run = p_cell.add_run("OBSERVACIÓN: ")
                run.font.bold = True
                p_cell.add_run(actividad["despues"]["observacion"])

                set_cell_border(
                    cell,
                    top={"sz": 12, "val": "single", "color": "000000"},
                    bottom={"sz": 12, "val": "single", "color": "000000"},
                    left={"sz": 12, "val": "single", "color": "000000"},
                    right={"sz": 12, "val": "single", "color": "000000"},
                )

                # Fila 3: Imágenes DESPUÉS
                cell_img = tabla_despues.rows[2].cells[0]
                if actividad["despues"].get("imagenes"):
                    for img_bytes in actividad["despues"]["imagenes"]:
                        add_image_to_cell(cell_img, img_bytes, width_inches=2.2)

                set_cell_border(
                    cell_img,
                    top={"sz": 12, "val": "single", "color": "000000"},
                    bottom={"sz": 12, "val": "single", "color": "000000"},
                    left={"sz": 12, "val": "single", "color": "000000"},
                    right={"sz": 12, "val": "single", "color": "000000"},
                )

        doc.add_paragraph()

    return doc


# ============= INTERFAZ DE STREAMLIT =============

# Sidebar para información general
with st.sidebar:
    st.header("📋 Información General")

    st.subheader("Datos de la Empresa")
    empresa_nombre_proyecto = st.text_input("Nombre del Proyecto", "LA RITA")
    empresa_fecha = st.text_input("Fecha del Informe", "NOVIEMBRE 2025")
    empresa_tecnico = st.text_input("Nombre del Técnico", "")
    empresa_ubicacion = st.text_input("Ubicación del Trabajo", "")

    st.subheader("Objetivo")
    empresa_objetivo = st.text_area(
        "Objetivo del Informe",
        "A continuación, se describe los trabajos de mantenimiento realizados en el Stard, tanques y cajas, así como las acciones realizadas para corregir las deficiencias con el fin de lograr mejor funcionamiento del sistema.",
        height=100,
    )

    st.subheader("Nota")
    empresa_nota = st.text_area(
        "Nota de Seguridad",
        "Antes de iniciar con cualquier tipo de proceso, nuestro personal técnico cuenta con todas las medidas de seguridad necesarias, ya que se encuentran expuestos a diferentes riesgos.",
        height=100,
    )

    st.markdown("---")
    st.subheader("Datos del Cliente")
    cliente_nombre = st.text_input("Nombre o Razón Social", "MANUELITA S.A")
    cliente_nit = st.text_input("NIT o C.C.", "891.300.241-9")
    cliente_direccion = st.text_input("Dirección", "kilometro 7 via palmira el cerrito")

# Área principal
tab1, tab2, tab3 = st.tabs(
    ["📝 Agregar Actividades", "👁️ Vista Previa", "💾 Generar Documento"]
)

with tab1:
    st.header("Agregar Actividades al Informe")

    # Inicializar session_state para actividades
    if "actividades" not in st.session_state:
        st.session_state.actividades = []

    # Selector de tipo de actividad
    col1, col2 = st.columns([2, 1])

    with col1:
        tipo_actividad = st.selectbox(
            "Tipo de Actividad",
            [
                "Cerramiento del Área de Trabajo",
                "Fumigación Alrededor de las Tapas",
                "Fumigación del Sendero",
                "Bypass de Entrada",
                "Bypass de Salida",
                "Limpieza de Pozo Séptico con Vactor",
                "Lavado de Filtros",
                "Aplicación de Tratamiento Séptico",
                "Limpieza de Tornillos/Compuertas",
                "Limpieza de Maleza",
                "Limpieza Alrededor de las Tapas",
                "Lavado de Tapas del Pozo",
                "Otra (personalizada)",
            ],
        )

    with col2:
        formato = st.radio("Formato", ["Solo Observación", "Antes/Después"])

    if tipo_actividad == "Otra (personalizada)":
        titulo_actividad = st.text_input("Nombre de la Actividad", "")
    else:
        titulo_actividad = tipo_actividad

    st.markdown("---")

    if formato == "Solo Observación":
        st.subheader("📸 Observación")
        observacion = st.text_area(
            "Descripción de la Observación",
            height=100,
            placeholder="Describa detalladamente lo observado durante esta actividad...",
        )

        imagenes = st.file_uploader(
            "Subir Fotografías",
            type=["png", "jpg", "jpeg"],
            accept_multiple_files=True,
            key="img_obs",
        )

        if st.button("➕ Agregar Actividad", type="primary", use_container_width=True):
            if titulo_actividad and observacion:
                # Procesar imágenes
                imgs_bytes = []
                if imagenes:
                    for img in imagenes:
                        imgs_bytes.append(img.read())

                actividad = {
                    "titulo": titulo_actividad,
                    "tipo": "solo_observacion",
                    "observacion": observacion,
                    "imagenes": imgs_bytes,
                }
                st.session_state.actividades.append(actividad)
                st.success(f"✅ Actividad '{titulo_actividad}' agregada correctamente!")
                st.rerun()
            else:
                st.error("⚠️ Por favor complete el título y la observación.")

    else:  # Antes/Después
        col_antes, col_despues = st.columns(2)

        with col_antes:
            st.subheader("📸 ANTES")
            obs_antes = st.text_area(
                "Observación ANTES",
                height=100,
                placeholder="Describa el estado inicial...",
                key="obs_antes",
            )
            imgs_antes = st.file_uploader(
                "Fotografías ANTES",
                type=["png", "jpg", "jpeg"],
                accept_multiple_files=True,
                key="imgs_antes",
            )

        with col_despues:
            st.subheader("📸 DESPUÉS")
            obs_despues = st.text_area(
                "Observación DESPUÉS",
                height=100,
                placeholder="Describa el estado final...",
                key="obs_despues",
            )
            imgs_despues = st.file_uploader(
                "Fotografías DESPUÉS",
                type=["png", "jpg", "jpeg"],
                accept_multiple_files=True,
                key="imgs_despues",
            )

        if st.button("➕ Agregar Actividad", type="primary", use_container_width=True):
            if titulo_actividad and (obs_antes or obs_despues):
                # Procesar imágenes ANTES
                antes_bytes = []
                if imgs_antes:
                    for img in imgs_antes:
                        antes_bytes.append(img.read())

                # Procesar imágenes DESPUÉS
                despues_bytes = []
                if imgs_despues:
                    for img in imgs_despues:
                        despues_bytes.append(img.read())

                actividad = {
                    "titulo": titulo_actividad,
                    "tipo": "antes_despues",
                    "antes": (
                        {"observacion": obs_antes, "imagenes": antes_bytes}
                        if obs_antes
                        else None
                    ),
                    "despues": (
                        {"observacion": obs_despues, "imagenes": despues_bytes}
                        if obs_despues
                        else None
                    ),
                }
                st.session_state.actividades.append(actividad)
                st.success(f"✅ Actividad '{titulo_actividad}' agregada correctamente!")
                st.rerun()
            else:
                st.error("⚠️ Por favor complete el título y al menos una observación.")

    st.markdown("---")

    # Lista de actividades agregadas
    if st.session_state.actividades:
        st.subheader(f"📋 Actividades Agregadas ({len(st.session_state.actividades)})")

        for idx, act in enumerate(st.session_state.actividades):
            with st.expander(f"{idx + 1}. {act['titulo']}"):
                if act["tipo"] == "solo_observacion":
                    st.write("**Observación:**", act["observacion"])
                    st.write(f"**Fotografías:** {len(act.get('imagenes', []))}")
                else:
                    if act.get("antes"):
                        st.write("**ANTES:**", act["antes"]["observacion"])
                        st.write(
                            f"  - Fotografías: {len(act['antes'].get('imagenes', []))}"
                        )
                    if act.get("despues"):
                        st.write("**DESPUÉS:**", act["despues"]["observacion"])
                        st.write(
                            f"  - Fotografías: {len(act['despues'].get('imagenes', []))}"
                        )

                if st.button(f"🗑️ Eliminar", key=f"del_{idx}"):
                    st.session_state.actividades.pop(idx)
                    st.rerun()

        if st.button("🗑️ Limpiar Todas las Actividades", type="secondary"):
            st.session_state.actividades = []
            st.rerun()

with tab2:
    st.header("👁️ Vista Previa del Informe")

    if not st.session_state.get("actividades"):
        st.info(
            "ℹ️ No hay actividades agregadas. Agregue actividades en la pestaña 'Agregar Actividades'."
        )
    else:
        st.success(f"✅ Total de actividades: {len(st.session_state.actividades)}")

        # Mostrar resumen
        st.subheader("Datos Generales")
        st.write(f"**Proyecto:** {empresa_nombre_proyecto}")
        st.write(f"**Fecha:** {empresa_fecha}")
        st.write(f"**Técnico:** {empresa_tecnico}")
        st.write(f"**Ubicación:** {empresa_ubicacion}")
        st.write(f"**Cliente:** {cliente_nombre}")

        st.markdown("---")
        st.subheader("Actividades Incluidas")

        for idx, act in enumerate(st.session_state.actividades):
            st.write(f"**{idx + 1}. {act['titulo']}**")
            if act["tipo"] == "solo_observacion":
                st.write(f"   - Tipo: Solo Observación")
                st.write(f"   - Fotografías: {len(act.get('imagenes', []))}")
            else:
                st.write(f"   - Tipo: Antes/Después")
                if act.get("antes"):
                    st.write(
                        f"   - Fotos ANTES: {len(act['antes'].get('imagenes', []))}"
                    )
                if act.get("despues"):
                    st.write(
                        f"   - Fotos DESPUÉS: {len(act['despues'].get('imagenes', []))}"
                    )

with tab3:
    st.header("💾 Generar Documento Word")

    if not st.session_state.get("actividades"):
        st.warning(
            "⚠️ Debe agregar al menos una actividad antes de generar el documento."
        )
    elif not all(
        [
            empresa_tecnico,
            empresa_ubicacion,
            cliente_nombre,
            cliente_nit,
            cliente_direccion,
        ]
    ):
        st.warning("⚠️ Por favor complete todos los datos en la barra lateral.")
    else:
        st.success("✅ Todo listo para generar el documento")

        st.write("**Resumen:**")
        st.write(f"- Actividades: {len(st.session_state.actividades)}")
        st.write(f"- Cliente: {cliente_nombre}")
        st.write(f"- Técnico: {empresa_tecnico}")

        if st.button(
            "📄 Generar Informe Técnico", type="primary", use_container_width=True
        ):
            with st.spinner("Generando documento..."):
                # Preparar datos
                datos_empresa = {
                    "nombre_proyecto": empresa_nombre_proyecto,
                    "fecha": empresa_fecha,
                    "tecnico": empresa_tecnico,
                    "ubicacion": empresa_ubicacion,
                    "objetivo": empresa_objetivo,
                    "nota": empresa_nota,
                }

                datos_cliente = {
                    "nombre": cliente_nombre,
                    "nit": cliente_nit,
                    "direccion": cliente_direccion,
                }

                # Crear documento
                doc = crear_documento_tecnico(
                    datos_empresa, datos_cliente, st.session_state.actividades
                )

                # Guardar en memoria
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)

                # Nombre del archivo
                fecha_str = datetime.now().strftime("%Y%m%d")
                nombre_archivo = f"INFORME_TECNICO_{empresa_nombre_proyecto.replace(' ', '_')}_{fecha_str}.docx"

                st.success("✅ ¡Documento generado exitosamente!")

                # Botón de descarga
                st.download_button(
                    label="⬇️ Descargar Informe Técnico",
                    data=doc_io,
                    file_name=nombre_archivo,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

# Footer
st.markdown("---")
st.markdown(
    """
    <div style='text-align: center; color: #666;'>
        <p>Generador de Informes Técnicos v1.0 | Desarrollado con Streamlit</p>
    </div>
    """,
    unsafe_allow_html=True,
)
